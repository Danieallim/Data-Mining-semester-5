# classification_bank_marketing_configured.py
# Dataset: Bank Marketing (CSV kamu). Latih 3 model, evaluasi lengkap, simpan grafik & artefak.
# Update: metrics CSV, CM raw+normalized, ROC/PR, top features + bar plot, threshold tuning, classification report,
#         optional DOCX auto (try/except).

import numpy as np
import pandas as pd
from pathlib import Path
import json
import matplotlib.pyplot as plt

from sklearn.model_selection import train_test_split
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.impute import SimpleImputer
from sklearn.metrics import (
    confusion_matrix, ConfusionMatrixDisplay,
    roc_auc_score, roc_curve, precision_recall_curve,
    average_precision_score, f1_score, precision_score, recall_score, accuracy_score,
    classification_report
)
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier, HistGradientBoostingClassifier

# =========================
# 0) KONFIGURASI
# =========================
# GANTI sesuai nama/path CSV kamu (boleh absolut)
CSV_PATH = "bank_marketing.csv"
TARGET = "y"
# "pr_auc" (disarankan untuk imbalance) atau "roc_auc"
SELECT_BY = "pr_auc"
TARGET_RECALL = 0.50                    # target recall untuk threshold tuning
TOPK_FEATURES = 15                      # berapa banyak fitur teratas disimpan/gambar
# contoh: 12000 (untuk laptop pelan). None = pakai semua data
SAMPLE_N = None

# =========================
# 1) LOAD & CLEAN
# =========================
df = pd.read_csv(CSV_PATH)

# Map target yes/no -> 1/0 (aman untuk variasi kapital/whitespace)
df[TARGET] = df[TARGET].astype(
    str).str.strip().str.lower().map({"yes": 1, "no": 0})

# Buang kolom leakage/ID (penting!)
for col in ["duration", "Id", "id"]:
    if col in df.columns:
        df = df.drop(columns=[col])

# (Opsional) sampling agar cepat
if SAMPLE_N is not None and SAMPLE_N < len(df):
    df = df.sample(n=SAMPLE_N, random_state=42, stratify=df[TARGET])

y = df[TARGET].values
X = df.drop(columns=[TARGET])

# =========================
# 2) PREPROCESSOR
# =========================
num_cols = X.select_dtypes(include=[np.number]).columns.tolist()
cat_cols = [c for c in X.columns if c not in num_cols]

num_tf = Pipeline([
    ("impute", SimpleImputer(strategy="median")),
    ("scale", StandardScaler())
])
cat_tf = Pipeline([
    ("impute", SimpleImputer(strategy="most_frequent")),
    ("onehot", OneHotEncoder(handle_unknown="ignore"))
])

pre = ColumnTransformer([
    ("num", num_tf, num_cols),
    ("cat", cat_tf, cat_cols)
])

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, stratify=y, random_state=42
)

# =========================
# 3) MODELS
# =========================
models = {
    "logreg": Pipeline([("pre", pre),
                        ("clf", LogisticRegression(max_iter=400, class_weight="balanced"))]),
    "rf": Pipeline([("pre", pre),
                    ("clf", RandomForestClassifier(
                        n_estimators=250, class_weight="balanced_subsample", random_state=42))]),
    "hgb": Pipeline([("pre", pre),
                     ("clf", HistGradientBoostingClassifier(random_state=42))]),
}

# =========================
# 4) TRAIN & EVAL (threshold 0.5)
# =========================
rows, trained = [], {}
for name, pipe in models.items():
    pipe.fit(X_train, y_train)
    trained[name] = pipe

    # Probabilitas untuk ROC/PR
    clf = pipe.named_steps["clf"]
    if hasattr(clf, "predict_proba"):
        proba = pipe.predict_proba(X_test)[:, 1]
    else:
        # fallback: beberapa estimator non-proba
        proba = pipe.predict(X_test).astype(float)

    pred = pipe.predict(X_test)

    rows.append({
        "model": name,
        "roc_auc": roc_auc_score(y_test, proba),
        "pr_auc": average_precision_score(y_test, proba),
        "f1": f1_score(y_test, pred),
        "recall": recall_score(y_test, pred),
        "precision": precision_score(y_test, pred, zero_division=0),
        "acc": accuracy_score(y_test, pred),
    })

res_df = pd.DataFrame(rows).sort_values(SELECT_BY, ascending=False)
print("\n=== Metrics (test set, threshold=0.5) ===")
print(res_df.to_string(index=False))

# =========================
# 5) OUTPUT DIRS
# =========================
out_dir = Path("outputs")
out_dir.mkdir(exist_ok=True)
fig_dir = Path("figs")
fig_dir.mkdir(exist_ok=True)

# simpan tabel metrik
(res_df).to_csv(out_dir / "metrics_all_models.csv", index=False)

# =========================
# 6) MODEL TERPILIH + FIGURES
# =========================
best_name = res_df.iloc[0]["model"]
best = trained[best_name]
clf_best = best.named_steps["clf"]

# Confusion Matrix (raw)
pred_best = best.predict(X_test)
cm = confusion_matrix(y_test, pred_best)
disp = ConfusionMatrixDisplay(cm)
disp.plot()
plt.title(f"Confusion Matrix — {best_name} (raw)")
plt.tight_layout()
plt.savefig(fig_dir / f"confmat_{best_name}_raw.png", dpi=160)
plt.close()

# Confusion Matrix (normalized by true labels)
ConfusionMatrixDisplay.from_estimator(best, X_test, y_test, normalize='true')
plt.title(f"Confusion Matrix — {best_name} (normalized)")
plt.tight_layout()
plt.savefig(fig_dir / f"confmat_{best_name}_norm.png", dpi=160)
plt.close()

# ROC & PR untuk model terbaik
if hasattr(clf_best, "predict_proba"):
    proba_best = best.predict_proba(X_test)[:, 1]
else:
    proba_best = best.predict(X_test).astype(float)

fpr, tpr, thr = roc_curve(y_test, proba_best)
plt.plot(fpr, tpr)
plt.plot([0, 1], [0, 1], '--')
plt.xlabel("FPR")
plt.ylabel("TPR")
plt.title(f"ROC — {best_name}")
plt.tight_layout()
plt.savefig(fig_dir / f"roc_{best_name}.png", dpi=160)
plt.close()

prec, rec, thr_pr = precision_recall_curve(y_test, proba_best)
plt.plot(rec, prec)
plt.xlabel("Recall")
plt.ylabel("Precision")
plt.title(f"PR — {best_name}")
plt.tight_layout()
plt.savefig(fig_dir / f"pr_{best_name}.png", dpi=160)
plt.close()

# Classification report (threshold 0.5)
report_txt = classification_report(y_test, pred_best, digits=3)
(out_dir /
 f"classification_report_{best_name}_t0.5.txt").write_text(report_txt, encoding="utf-8")

# =========================
# 7) TOP FEATURES
# =========================
# Ambil nama fitur: numerik (asli) + one-hot untuk kategorikal
ohe = best.named_steps["pre"].named_transformers_["cat"].named_steps["onehot"]
cat_feat_names = ohe.get_feature_names_out(
    cat_cols) if len(cat_cols) else np.array([])
feat_names = np.r_[num_cols, cat_feat_names]

top_df = None
if hasattr(clf_best, "feature_importances_"):
    vals = clf_best.feature_importances_
    top_df = pd.DataFrame({"feature": feat_names, "value": vals})\
        .sort_values("value", ascending=False).head(TOPK_FEATURES)
    top_df.to_csv(out_dir / f"top_features_{best_name}.csv", index=False)
elif hasattr(clf_best, "coef_"):
    vals = np.abs(clf_best.coef_.ravel())
    idx = np.argsort(vals)[::-1][:TOPK_FEATURES]
    top_df = pd.DataFrame({"feature": feat_names[idx], "value": vals[idx]})
    top_df.to_csv(out_dir / f"top_features_{best_name}.csv", index=False)

# Bar plot top features (jika tersedia)
if top_df is not None and len(top_df) > 0:
    plt.figure(figsize=(8, max(4, 0.35*len(top_df))))
    plt.barh(top_df["feature"][::-1], top_df["value"][::-1])
    plt.xlabel("Importance / |coef|")
    plt.title(f"Top {len(top_df)} Features — {best_name}")
    plt.tight_layout()
    plt.savefig(fig_dir / f"top_features_{best_name}.png", dpi=160)
    plt.close()

# =========================
# 8) THRESHOLD TUNING (target recall)
# =========================
# pilih threshold yang mencapai recall >= TARGET_RECALL dengan precision tertinggi
thr_candidates = []
prec_all, rec_all, thr_all = precision_recall_curve(y_test, proba_best)
# pad thr agar sama panjang dengan prec/rec bila perlu
thr_pad = np.r_[thr_all, 1.0] if len(thr_all) == len(prec_all) - 1 else thr_all

for p, r, t in zip(prec_all, rec_all, thr_pad):
    if r >= TARGET_RECALL:
        thr_candidates.append((p, r, t))
if thr_candidates:
    # precision max di recall >= target
    p_sel, r_sel, t_sel = max(thr_candidates, key=lambda x: x[0])
    y_sel = (proba_best >= t_sel).astype(int)
    cm_sel = confusion_matrix(y_test, y_sel)
    metrics_sel = {
        "threshold": float(t_sel),
        "precision": float(precision_score(y_test, y_sel, zero_division=0)),
        "recall": float(recall_score(y_test, y_sel)),
        "f1": float(f1_score(y_test, y_sel)),
        "acc": float(accuracy_score(y_test, y_sel)),
        "roc_auc": float(roc_auc_score(y_test, proba_best)),
        "pr_auc": float(average_precision_score(y_test, proba_best)),
        "cm": cm_sel.tolist()
    }
    (out_dir / f"threshold_metrics_{best_name}.json").write_text(
        json.dumps(metrics_sel, indent=2), encoding="utf-8")

    # simpan CM di threshold terpilih
    disp = ConfusionMatrixDisplay(cm_sel)
    disp.plot()
    plt.title(f"Confusion Matrix — {best_name} (threshold={t_sel:.3f})")
    plt.tight_layout()
    plt.savefig(fig_dir / f"confmat_{best_name}_thr{t_sel:.3f}.png", dpi=160)
    plt.close()

    # classification report di threshold terpilih
    rep_sel = classification_report(y_test, y_sel, digits=3)
    (out_dir / f"classification_report_{best_name}_thr{t_sel:.3f}.txt").write_text(
        rep_sel, encoding="utf-8")

    print(f"\n=== Threshold tuning (target recall {TARGET_RECALL:.2f}) ===")
    print(
        f"chosen threshold = {t_sel:.3f} | precision {metrics_sel['precision']:.3f} | recall {metrics_sel['recall']:.3f} | f1 {metrics_sel['f1']:.3f}")
else:
    print(
        f"\n[INFO] Tidak ada threshold yang mencapai recall >= {TARGET_RECALL:.2f}. Pertimbangkan rebalancing/SMOTE atau target recall lebih rendah.")

# =========================
# 9) OPTIONAL: AUTO-DOCX (jika python-docx terpasang)
# =========================
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading("Bank Marketing — Term Deposit Classification", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        "Author: AA\nFinal Project — Findings, Insights & Model Selection")
    p.paragraph_format.space_after = Pt(6)

    # Summary singkat
    doc.add_heading("1) Executive Summary (ringkas)", level=1)
    doc.add_paragraph(
        f"Model terbaik berdasarkan {SELECT_BY.upper()} adalah {best_name.upper()}. "
        "Kami menyarankan threshold tuning berbasis biaya dan penggunaan skor untuk memprioritaskan follow-up."
    )

    # Tabel metrik semua model
    doc.add_heading("2) Results (all models)", level=1)
    t = doc.add_table(rows=1, cols=7)
    for i, h in enumerate(["Model", "ROC-AUC", "PR-AUC", "F1", "Recall", "Precision", "Accuracy"]):
        t.rows[0].cells[i].text = h
    for _, r in res_df.iterrows():
        row = t.add_row().cells
        row[0].text = str(r["model"])
        row[1].text = f'{r["roc_auc"]:.3f}'
        row[2].text = f'{r["pr_auc"]:.3f}'
        row[3].text = f'{r["f1"]:.3f}'
        row[4].text = f'{r["recall"]:.3f}'
        row[5].text = f'{r["precision"]:.3f}'
        row[6].text = f'{r["acc"]:.3f}'

    # Sisipkan gambar
    doc.add_heading("3) Evaluation Figures", level=1)
    for pth in [f"confmat_{best_name}_raw.png", f"confmat_{best_name}_norm.png",
                f"roc_{best_name}.png", f"pr_{best_name}.png",
                f"top_features_{best_name}.png"]:
        p_img = fig_dir / pth
        if p_img.exists():
            doc.add_paragraph(pth.replace("_", " ").replace(".png", ""))
            doc.add_picture(str(p_img), width=Inches(5.5))

    # Save DOCX
    out_docx = out_dir / "Auto_Report_Bank_Marketing.docx"
    doc.save(out_docx)
    print(f"\n[DOCX] Generated: {out_docx}")
except Exception as e:
    print("\n[DOCX] Skip (python-docx tidak terpasang atau ada error ringan).")
